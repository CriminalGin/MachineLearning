import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from numpy import shape, mean, round
from sklearn import datasets, linear_model
from sklearn.preprocessing import StandardScaler
import seaborn
from docx import Document
from docx.shared import Inches, Pt
seaborn.set()
document = Document('Assignment1.docx')

feature_name = ['symboling', 'normalized-losses', 'make', 'fuel-type', 'aspiration', 'num-of-doors', 'body-style', \
                         'drive-wheels', 'engine-location', 'wheel-base', 'length', 'width', 'height', 'curb-weight', 'engine-type',\
                         'num-of-cylinders', 'engine-size', 'fuel-system', 'bore', 'stroke', 'compression-ratio', \
                         'horsepower', 'peak-rpm', 'city-mpg', 'highway-mpg', 'price']


# 3.2.1
df = pd.read_csv('E:\Document Files\Learning\Fund of ML\PythonCode\Prog1\imports-85.data', \
                 header=None, names = feature_name, na_values='?')

# 3.2.2
df = df.dropna()

numpy_df = df.values    # Use its value

# 3.2.3
numpy_df_train = numpy_df[:int(round(shape(numpy_df)[0] * 0.8))]
numpy_df_test = numpy_df[:int(shape(numpy_df)[0] - round(shape(numpy_df)[0] * 0.8))]

model = linear_model.LinearRegression()
numpy_df_train_x = np.expand_dims(numpy_df_train[:,16], axis=1)
numpy_df_train_y = numpy_df_train[:,25]

numpy_df_test_x = np.expand_dims(numpy_df_test[:,16], axis=1)
numpy_df_test_y = numpy_df_test[:,25]
model.fit(numpy_df_train_x, numpy_df_train_y)    # The x-axis matrice must be 2D!

numpy_df_test_y_predict = model.predict(np.expand_dims(numpy_df_test[:,16],axis=1))

plt.scatter(numpy_df_test_x, numpy_df_test_y, c = 'b',  marker='o')
plt.scatter(numpy_df_test_x, numpy_df_test_y_predict, c = 'r',  marker='*')
plt.xlabel('Engine size')
plt.ylabel('Price')
plt.title('Linear regression on clean data')
plt.savefig('3.2.3.png')
plt.show()
document.add_paragraph('3.2.3', style = 'Normal')
document.add_picture('3.2.3.png', width=Inches(6.2))

print 'Price prediction for engine size equals to 175 is: %f' %(model.predict(175))
document.add_paragraph('Price prediction for engine size equals to 175 is: %f' %(model.predict(175)), style = 'Normal')



# 3.2.4
X_scaler = StandardScaler()
numpy_df_train_x_scaled = X_scaler.fit_transform(numpy_df_train_x)
numpy_df_train_y_scaled = X_scaler.fit_transform(numpy_df_train_y)
numpy_df_test_x_scaled = X_scaler.transform(numpy_df_test_x)
numpy_df_test_y_scaled = X_scaler.transform(numpy_df_test_y)

model_standard = linear_model.LinearRegression()
model_standard.fit(numpy_df_train_x_scaled, numpy_df_train_y_scaled)
numpy_df_test_y_scaled_predict = model_standard.predict(numpy_df_test_x_scaled)

plt.scatter(numpy_df_test_x_scaled, numpy_df_test_y_scaled, c = 'b',  marker='o')
plt.scatter(numpy_df_test_x_scaled, numpy_df_test_y_scaled_predict, c = 'r',  marker='*')
plt.xlabel('Standardized Engine-size')
plt.ylabel('Standardized Price')
plt.title('Linear regression on Standardized data')
plt.savefig('3.2.4.png')
plt.show()
document.add_paragraph('3.2.4', style = 'Normal')
document.add_picture('3.2.4.png', width=Inches(6.2))


# document.save('Assignment1.docx')
