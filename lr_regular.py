import numpy as np
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
from sklearn.linear_model import Ridge
from sklearn.preprocessing import PolynomialFeatures
from docx import Document
from docx.shared import Inches, Pt

document = Document('Assignment1.docx')

X_train = [[6], [8], [10], [14], [18]]
y_train = [[7.5], [9.1], [13.2], [17.5], [19.3]]

X_test = [[6], [8], [11], [16]]
y_test = [[8.3], [12.5], [15.4], [18.6]]

# 3.3.1
lr_model = LinearRegression()
lr_model.fit(X_train, y_train)

print 'y1= %f + %f x' %(lr_model.intercept_, lr_model.coef_)

document.add_paragraph('3.3.1 LR regression on polynomial data', style = 'Heading 2')
document.add_paragraph('y1= %f + %f x' %(lr_model.intercept_, lr_model.coef_), style = 'Normal')

lr_score = lr_model.score(X_test, y_test)
print 'Linear regression (order 1) model score is: ',
print lr_score
document.add_paragraph('Linear regression (order 1) model score is: %f' %(lr_score), style = 'Normal')

xx = np.linspace(0, 26, 100)
yy = lr_model.predict(xx.reshape(xx.shape[0], 1))

plt.plot(xx, yy)
plt.scatter(X_test, y_test, c = 'r',  marker='o')
plt.title('Linear regression (order 1) result')
plt.savefig('3.3.1.png')
plt.show()
document.add_picture('3.3.1.png', width=Inches(6))

# 3.3.2
poly = PolynomialFeatures(degree=5)
X_train_poly = poly.fit_transform(X_train)
X_test_poly = poly.transform(X_test)
poly_lr_model = LinearRegression()
poly_lr_model.fit(X_train_poly, y_train)

document.add_paragraph('3.3.2 Polynomial regression on training data', style = 'Heading 2')

print 'y2= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(poly_lr_model.intercept_, poly_lr_model.coef_[0, 1], poly_lr_model.coef_[0, 2], poly_lr_model.coef_[0, 3], \
        poly_lr_model.coef_[0, 4], poly_lr_model.coef_[0, 5])

document.add_paragraph('y2= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(poly_lr_model.intercept_, poly_lr_model.coef_[0, 1], poly_lr_model.coef_[0, 2], poly_lr_model.coef_[0, 3], \
        poly_lr_model.coef_[0, 4], poly_lr_model.coef_[0, 5]))

xx_poly = poly.transform(xx.reshape(xx.shape[0], 1))
yy_poly = poly_lr_model.predict(xx_poly)

poly_lr_score = poly_lr_model.score(X_test_poly, y_test)

print 'Linear regression (order 5) score is: ',
print poly_lr_score

document.add_paragraph('Linear regression (order 5) score is: %f' %(poly_lr_score))


plt.plot(xx, yy_poly)
plt.scatter(X_test, y_test, c = 'r',  marker='o')
plt.title('Linear regression (order 5) result')
plt.savefig('3.3.2.png')
plt.show()

document.add_picture('3.3.2.png', width=Inches(6))

# 3.3.3
ridge_model = Ridge(alpha=1, normalize=False)
ridge_model.fit(X_train_poly, y_train)
yy_ridge = ridge_model.predict(xx_poly) # get prediction for xx_poly
ridge_score = ridge_model.score(X_test_poly, y_test)
print 'Ridge regression (order 5) score is: ',
print ridge_score
document.add_paragraph('3.3.3 Ridge Regression', style = 'Heading 2')
document.add_paragraph('Ridge regression (order 5) score is: %f' %(ridge_score))

print 'y3= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(ridge_model.intercept_, ridge_model.coef_[0, 1], ridge_model.coef_[0, 2], ridge_model.coef_[0, 3], \
        ridge_model.coef_[0, 4], ridge_model.coef_[0, 5])

document.add_paragraph('y3= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(ridge_model.intercept_, ridge_model.coef_[0, 1], ridge_model.coef_[0, 2], ridge_model.coef_[0, 3], \
        ridge_model.coef_[0, 4], ridge_model.coef_[0, 5]))

plt.plot(xx, yy_ridge)
plt.scatter(X_test, y_test, c = 'r',  marker='o')
plt.title('Ridge regression (order 5) result')
plt.savefig('3.3.3.png')
plt.show()
document.add_picture('3.3.3.png', width=Inches(6))

# 3.3.4
document.add_paragraph('3.3.4 Comparisons', style = 'Heading 2')
document.add_paragraph('The model with the highest score is: the ridge regression(order 5)')
document.add_paragraph('Ridge model can prevent over-fitting: yes')

ridge_model = Ridge(alpha=0, normalize=False)
ridge_model.fit(X_train_poly, y_train)
yy_ridge = ridge_model.predict(xx_poly) # get prediction for xx_poly
'''
print 'Linear regression (order 5) y2= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(poly_lr_model.intercept_, poly_lr_model.coef_[0, 0], poly_lr_model.coef_[0, 1], poly_lr_model.coef_[0, 2], \
        poly_lr_model.coef_[0, 3], poly_lr_model.coef_[0, 4])

print 'Ridge regression (order 5) alpha = 0 y3= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(ridge_model.intercept_, ridge_model.coef_[0, 0], ridge_model.coef_[0, 1], ridge_model.coef_[0, 2], \
        ridge_model.coef_[0, 3], ridge_model.coef_[0, 4])
'''
document.add_paragraph('Ridge model is nearly equivalent to LR model (order 5) if alpha=0: yes')

ridge_model = Ridge(alpha=5, normalize=False)
ridge_model.fit(X_train_poly, y_train)
yy_ridge = ridge_model.predict(xx_poly) # get prediction for xx_poly

print 'y3= %f + %f x + %f x*x + %f x*x*x + %f x*x*x*x + %f x*x*x*x*x' \
      %(ridge_model.intercept_, ridge_model.coef_[0, 1], ridge_model.coef_[0, 2], ridge_model.coef_[0, 3], \
        ridge_model.coef_[0, 4], ridge_model.coef_[0, 5])

document.add_paragraph('A larger alpha results in a larger coefficient for x*x*x*x*x: no')

# Optional
poly_lr_score_max = 0
ridge_score_max = 0
for i in range(1000):
    i = i / 100
    poly = PolynomialFeatures(degree=i)
    X_train_poly = poly.fit_transform(X_train)
    X_test_poly = poly.transform(X_test)
    poly_lr_model = LinearRegression()
    poly_lr_model.fit(X_train_poly, y_train)

    xx_poly = poly.transform(xx.reshape(xx.shape[0], 1))
    yy_poly = poly_lr_model.predict(xx_poly)

    poly_lr_score = poly_lr_model.score(X_test_poly, y_test)
    if poly_lr_score > poly_lr_score_max:
        poly_lr_score_max = poly_lr_score
        poly_lr_score_max_i = i

    ridge_model = Ridge(alpha=1, normalize=False)
    ridge_model.fit(X_train_poly, y_train)
    yy_ridge = ridge_model.predict(xx_poly) # get prediction for xx_poly
    ridge_score = ridge_model.score(X_test_poly, y_test)
    if ridge_score > ridge_score_max:
        ridge_score_max = ridge_score
        ridge_score_max_i = i
document.add_paragraph('The best degree parameter for polynomial regression and Ridge regression with (alpha=1) is: %f' \
                       %(ridge_score_max_i))

document.save('Assignment1.docx')




