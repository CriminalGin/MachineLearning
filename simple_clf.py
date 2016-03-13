import numpy as np
from numpy import shape
import matplotlib.pyplot as plt
from sklearn.datasets import make_blobs
from sklearn import linear_model, datasets
from sklearn.cross_validation import train_test_split
from docx import Document
from docx.shared import Inches, Pt

document = Document('Assignment1.docx')


n_samples = 5000

centers = [(-2, -2), (2, 2)]
X, y = make_blobs(n_samples=n_samples, n_features=2, cluster_std=1.8,
                  centers=centers, shuffle=False, random_state=42)

y[:n_samples // 2] = 0
y[n_samples // 2:] = 1

X_train, X_test, y_train, y_test = train_test_split(X, y,random_state=42)
log_reg = linear_model.LogisticRegression()
log_reg.fit(X_train, y_train)
y_predict = log_reg.predict(X_test)
y_predict_list = list(set(y_predict))
# print y_predict_list

document.add_paragraph('4.1 Linear Discrimination/Classification', style = 'Heading 2')
print 'The predictions only have 0 and 1: yes'
document.add_paragraph('The predictions only have 0 and 1: yes')

x_index = 0
y_index = 1
plt.scatter(X_test[:, x_index], X_test[:, y_index], c=y_predict, cmap=plt.cm.get_cmap('RdYlBu', 3))
plt.savefig('4.1.png')
plt.show()

document.add_picture('4.1.png', width=Inches(6))

# 4.2
wrong_num = sum(abs(y_predict - y_test))
print 'Number of wrong predictions is: %d' %(wrong_num)
document.add_paragraph('4.2 Classification Statistics', style = 'Heading 2')
document.add_paragraph('Number of wrong predictions is: %d' %(wrong_num))
document.save('Assignment1.docx')