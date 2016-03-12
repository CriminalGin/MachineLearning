import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from numpy import dot, transpose, hstack, ones, shape
from numpy.linalg import inv
from sklearn import datasets, linear_model
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches, Pt
import seaborn
seaborn.set()

# 3.2.5

# 3.2.5.1
document = Document('Assignment1.docx')

feature_name = ['symboling', 'normalized-losses', 'make', 'fuel-type', 'aspiration', 'num-of-doors', 'body-style', \
                         'drive-wheels', 'engine-location', 'wheel-base', 'length', 'width', 'height', 'curb-weight', 'engine-type',\
                         'num-of-cylinders', 'engine-size', 'fuel-system', 'bore', 'stroke', 'compression-ratio', \
                         'horsepower', 'peak-rpm', 'city-mpg', 'highway-mpg', 'price']
df = pd.read_csv('imports-85.data', \
                 header=None, names = feature_name, na_values='?')
df = df.dropna()

numpy_df = df.values

numpy_df_engine_size = np.expand_dims(numpy_df[:,16], axis=1)
numpy_df_peak_rpm = np.expand_dims(numpy_df[:,22], axis=1)
numpy_df_price = numpy_df[:,25]

engine_size_scaler = StandardScaler()
numpy_df_engine_size_scaled = engine_size_scaler.fit_transform(numpy_df_engine_size)
# print numpy_df_engine_size_scaled

peak_rpm_scaler = StandardScaler()
numpy_df_peak_rpm_scaled = peak_rpm_scaler.fit_transform(numpy_df_peak_rpm)
# print numpy_df_peak_rpm_scaled

price_scaler = StandardScaler()
numpy_df_price_scaled = price_scaler.fit_transform(numpy_df_price)
# print numpy_df_price_scaled

# 3.2.5.2
feature_matrix = hstack([ones((shape(numpy_df_engine_size_scaled)[0], 1)), \
                 numpy_df_engine_size_scaled, numpy_df_peak_rpm_scaled])
parameter_theta = dot(dot(inv(dot(transpose(feature_matrix), feature_matrix)), transpose(feature_matrix)), numpy_df_price_scaled)
print 'Parameter theta calculated by normal equation: ',
print parameter_theta

document.add_paragraph('3.2.5', style = 'Normal')
p = document.add_paragraph('Parameter theta calculated by normal equation: [ %f %f %f]' %(parameter_theta[0], parameter_theta[1],  parameter_theta[2]))

# 3.2.5.3
sum = [0, 0, 0]
for i in range(100):
    i = i / 100
    for j in range(10):
        clf = linear_model.SGDRegressor(alpha = i)

        feature_matrix = hstack([numpy_df_engine_size_scaled, numpy_df_peak_rpm_scaled])
        clf.fit(feature_matrix, numpy_df_price_scaled)
        sum = hstack([clf.intercept_, clf.coef_]) + sum
average = sum / 1000
print 'Parameter theta calculated by SGD: ',
print average
p = document.add_paragraph('Parameter theta calculated by SGD: [ %f %f %f]' %(average[0], average[1],  average[2]))

document.save('Assignment1.docx')




