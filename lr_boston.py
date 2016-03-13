import matplotlib.pyplot as plt
import numpy as np
import matplotlib.pyplot as plt
from numpy import shape, mean
from sklearn import datasets, linear_model
import seaborn; seaborn.set()
from docx import Document
from docx.shared import Inches, Pt

document = Document()

document.add_heading('Program Assignment', 0)
document.add_paragraph('1155080708 JIANG Zhijian', style='Heading 3')

# Add a style
styleTitle = document.styles['Title']
font = styleTitle.font
font.name = 'Times New Roman'
font.size = Pt(18)

styleResult = document.styles['Normal']
font = styleResult.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Load the diabetes dataset
boston = datasets.load_boston()

# which feature
i_feature = 0
# Get the feature name
feature_name = boston.feature_names[i_feature]

# Use only one feature
diabetes_X = boston.data[:, np.newaxis, i_feature]

# Split the data into training/testing sets
boston_X_train = diabetes_X[:-20]
boston_X_test = diabetes_X[-20:]

# Split the targets into training/testing sets
boston_y_train = boston.target[:-20]
boston_y_test = boston.target[-20:]

# Create linear regression object
model = linear_model.LinearRegression()

# Train the model using the training sets
model.fit(boston_X_train, boston_y_train)

# Explained variance score: score=1 is perfect prediction
model_score = model.score(boston_X_test, boston_y_test)

# 3.1.1
print 'Number of features in the Boston dataset is: %d' %(shape(boston.data)[1])
print 'Number of samples in the Boston dataset is: %d' %(shape(boston.data)[0])

document.add_paragraph('3.1.1 Get n_feature and n_samples', style = 'Heading 2')
document.add_paragraph('Number of features in the Boston dataset is: %d' %(shape(boston.data)[1]), style = 'Normal')
document.add_paragraph('Number of samples in the Boston dataset is: %d' %(shape(boston.data)[0]), style = 'Normal')

# 3.1.2
model_score_max = -100
for i_feature in range(shape(boston.data)[1]):
    feature_name = boston.feature_names[i_feature]
    diabetes_X = boston.data[:, np.newaxis, i_feature]
    boston_X_train = diabetes_X[:-20]
    boston_X_test = diabetes_X[-20:]
    boston_y_train = boston.target[:-20]
    boston_y_test = boston.target[-20:]
    model = linear_model.LinearRegression()
    model.fit(boston_X_train, boston_y_train)
    model_score = model.score(boston_X_test, boston_y_test)
    if model_score > model_score_max:
        model_score_max = model_score
        model_score_max_name = feature_name
        model_score_max_position = i_feature
        best_model = model
model = best_model  # important!!
print 'Best fitted feature name is: %s' %(model_score_max_name)
print 'Best fitted model score is: ',
print model_score_max

document.add_paragraph('3.1.2 Find best fitted feature', style = 'Heading 2')
document.add_paragraph('Best fitted feature name is: %s' %(model_score_max_name), style = 'Normal')
document.add_paragraph('Best fitted model score is: %lf' %(model_score_max), style = 'Normal')

# 3.1.3
diabetes_X = boston.data[:, np.newaxis, model_score_max_position]
boston_X_test = diabetes_X[-20:]
boston_Y_test = boston.target[-20:]
boston_X_test_Y_predict = model.predict(boston_X_test)
loss = mean((boston_X_test_Y_predict - boston_Y_test) ** 2)
print 'Value of the loss function for the best fitted model is: %f' %(loss)

document.add_paragraph('3.1.3 Caculate the loss function', style = 'Heading 2')
document.add_paragraph('Value of the loss function for the best fitted model is: %f' %(loss), style = 'Normal')

# 3.1.4
plt.scatter(boston_X_test, boston_Y_test, c = 'b',  marker='o')
plt.scatter(boston_X_test, boston_X_test_Y_predict, c = 'r',  marker='*')
plt.xlabel(model_score_max_name)
plt.ylabel('Boston House Prices')
plt.savefig('3.1.4.png')
plt.show()
document.add_paragraph('3.1.4 Plot the predictions and test data', style = 'Heading 2')
document.add_picture('3.1.4.png', width=Inches(6))

document.save('Assignment1.docx')